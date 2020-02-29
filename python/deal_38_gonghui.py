from docx import Document
from openpyxl import Workbook

#  处理单选
schoice=Document('fkS.docx')
paras=schoice.parargraphs

swb=Workbook()
ws=swb.active

for i in range(0,len(paras),5):
    q=paras[i].text[2:-1]
    ans=paras[i].text[-1].upper()
    qs=[q,'单选',ans]
    qs.append(paras[i+1].text[2:])
    qs.append(paras[i+2].text[2:])
    qs.append(paras[i+3].text[2:])
    qs.append(paras[i+4].text[2:])
    ws.append(qs)
swb.save('fkS.xlsx')

# 处理多选题
mchoice=Document('fkM.docx')
mpas=mchoice.paragraphs
mwb=Workbook()
mws=mwb.active
for i in range(0,len(mpas),5):
    s=mpas[i].text.split('.')[1].split('-')
    q=s[0]
    ans=s[1].strip(' ').upper()
    qs=[q,'多选',ans]
    qs.append(mpas[i+1].text[2:])
    qs.append(mpas[i+2].text[2:])
    qs.append(mpas[i+3].text[2:])
    qs.append(mpas[i+4].text[2:])
    mws.append(qs)

mwb.save('fkM.xlsx')

# 处理判断题
jchoice=Document('fkJ.docx')
jpas=jchoice.paragraphs
jwb=Workbook()
jws=jwb.active
for i in range(0,len(jpas)):
    s=jpas[i].text.split('.')[1]
    q=s[0:-1]
    ans=s[-1].upper()
    qs=[q,'判断',ans]
    jws.append(qs)

jwb.save('fkJ.xlsx')